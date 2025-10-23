import os
from pathlib import Path
from typing import Dict, Optional
import PyPDF2
import docx


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.md'}
    
    def __init__(self):
        pass
    
    def load_document(self, file_path: str) -> Dict[str, str]:
        path = Path(file_path)
        
        # Validation
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {path.suffix}. "
                f"Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        # Extract text based on file type
        extension = path.suffix.lower()
        
        try:
            if extension == '.txt' or extension == '.md':
                text = self._load_text(file_path)
            elif extension == '.pdf':
                text = self._load_pdf(file_path)
            elif extension == '.docx':
                text = self._load_docx(file_path)
            else:
                raise ValueError(f"Unsupported extension: {extension}")
            
            # Validate extracted text
            if not text or len(text.strip()) == 0:
                raise ValueError(f"No text extracted from {path.name}")
            
            return {
                'text': text,
                'filename': path.name,
                'file_type': extension[1:]  # Remove the dot
            }
            
        except Exception as e:
            raise RuntimeError(f"Error loading {path.name}: {str(e)}")
    
    def _load_text(self, file_path: str) -> str:
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode file with any known encoding")
    
    def _load_pdf(self, file_path: str) -> str:
        text = []
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    raise ValueError("PDF is encrypted and cannot be read")
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text.append(page_text)
                    except Exception as e:
                        # Continue with other pages if one fails
                        print(f"Warning: Could not extract page {page_num + 1}: {e}")
                        continue
                
                if not text:
                    raise ValueError("No text could be extracted from PDF")
                
                return '\n\n'.join(text)
                
        except PyPDF2.errors.PdfReadError as e:
            raise ValueError(f"Invalid or corrupted PDF file: {e}")
    
    def _load_docx(self, file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables (important for structured data)
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
            
            # Combine all text
            all_text = paragraphs + table_text
            
            if not all_text:
                raise ValueError("No text found in document")
            
            return '\n\n'.join(all_text)
            
        except Exception as e:
            raise ValueError(f"Error reading Word document: {e}")
    
    @staticmethod
    def is_supported(file_path: str) -> bool:
        """Check if file type is supported"""
        return Path(file_path).suffix.lower() in DocumentLoader.SUPPORTED_EXTENSIONS


# Example usage (for testing)
if __name__ == "__main__":
    loader = DocumentLoader()
    
    # Test with a sample file
    try:
        doc = loader.load_document("sample.txt")
        print(f"Loaded {doc['filename']}")
        print(f"Text length: {len(doc['text'])} characters")
        print(f"First 200 chars: {doc['text'][:200]}...")
    except Exception as e:
        print(f"Error: {e}")